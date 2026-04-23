#!/usr/bin/env python3
"""Generate a meeting-ready document for the Juncheng update.

This script creates:
1. A Markdown version for easy editing.
2. A DOCX version that can be uploaded to Google Docs directly.
"""

from __future__ import annotations

import csv
import math
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/realtmxi/Desktop/NSDI2027_RouteWise")
OUTPUT_DIR = ROOT / "MEMORY"
DOCX_OUTPUT = OUTPUT_DIR / "juncheng_meeting_update_20260414.docx"
MD_OUTPUT = OUTPUT_DIR / "juncheng_meeting_update_20260414.md"

LATEST_V2_RUN = (
    ROOT
    / "experiment/results/v2_qwen3_90min/run_20260414_183856/evaluation_log.csv"
)
QWEN_COUNTERFACTUAL = (
    ROOT / "experiment/results/counterfactual_qwen3_no_wandb_v2/summary.csv"
)
MINIMAX_COUNTERFACTUAL = (
    ROOT / "experiment/results/counterfactual_minimax_no_inceptron_v2/summary.csv"
)


@dataclass(frozen=True)
class FigureSpec:
    """Figure metadata."""

    title: str
    path: Path
    caption: str
    width_in: float = 6.5


def _pct(values: list[float], q: float) -> float:
    if not values:
        return math.nan
    if len(values) == 1:
        return values[0]
    k = (len(values) - 1) * q
    f = math.floor(k)
    c = math.ceil(k)
    if f == c:
        return values[int(k)]
    return values[f] * (c - k) + values[c] * (k - f)


def _load_online_summary(csv_path: Path) -> list[dict[str, str]]:
    rows = list(csv.DictReader(csv_path.open()))
    policies = [
        "cheapest_fixed",
        "sort_latency",
        "smart_hedge",
        "v2_p50_hedge",
        "lp_mix",
        "openrouter_auto",
    ]
    summary = []
    for policy in policies:
        policy_rows = [row for row in rows if row["policy"] == policy]
        success_rows = [row for row in policy_rows if row["status"] == "success"]
        ttfts = sorted(
            float(row["ttft_ms"]) for row in success_rows if float(row["ttft_ms"]) >= 0
        )
        total_cost = sum(float(row["cost_usd"]) for row in policy_rows)
        slo_rate = (
            sum(row["slo_violated"].lower() == "true" for row in policy_rows)
            / len(policy_rows)
        )
        summary.append(
            {
                "policy": policy,
                "n": str(len(policy_rows)),
                "slo": f"{slo_rate * 100:.2f}%",
                "p50": f"{_pct(ttfts, 0.5):.0f}",
                "p99": f"{_pct(ttfts, 0.99):.0f}",
                "total_cost": f"${total_cost:.4f}",
            }
        )
    return summary


def _load_counterfactual_rows(csv_path: Path) -> list[dict[str, str]]:
    rows = list(csv.DictReader(csv_path.open()))
    selected = []
    for row in rows:
        if row["policy"] not in {
            "fastest_fixed",
            "smart_hedge",
            "lp_mix",
            "oracle_per_window",
            "cheapest_fixed",
        }:
            continue
        selected.append(
            {
                "policy": row["policy"],
                "slo": f"{float(row['slo_violation_rate_mean']) * 100:.2f}%",
                "slo_ci": (
                    f"[{float(row['slo_violation_rate_ci_lo']) * 100:.2f}, "
                    f"{float(row['slo_violation_rate_ci_hi']) * 100:.2f}]%"
                ),
                "p99": f"{float(row['p99_ms_mean']):.0f}",
                "cost": f"${float(row['mean_cost_mean']):.6f}",
                "hedge_rate": f"{float(row['hedge_rate_mean']) * 100:.1f}%",
            }
        )
    return selected


def _get_figures() -> list[FigureSpec]:
    return [
        FigureSpec(
            title="Figure 1. Production ACF Summary",
            path=ROOT
            / "experiment/results/acf_analysis_production/e2e_v2_gap_heatmap.png",
            caption=(
                "Each cell is P50 ACF(1) minus P99 ACF(1). Positive values mean "
                "P50 is more temporally stable than P99."
            ),
        ),
        FigureSpec(
            title="Figure 2. Production Full ACF Curves",
            path=ROOT
            / "experiment/results/acf_analysis_production/e2e_v2_acf_15min_min5.png",
            caption=(
                "The full ACF curves show that P50 typically stays above P99 "
                "across lags, not only at lag-1."
            ),
        ),
        FigureSpec(
            title="Figure 3. Live Qwen3 Lag-1 Scatter",
            path=ROOT
            / "experiment/results/phase5_qwen3_7d_clean/run_20260410_171625/analysis/autocorrelation/lag1_scatter_p50_vs_p99.png",
            caption=(
                "Supportive live-run evidence: P50 points cluster more tightly "
                "around the diagonal than P99."
            ),
        ),
        FigureSpec(
            title="Figure 4. Qwen3 Provider Shares Over Time",
            path=ROOT
            / "experiment/results/phase5_qwen3_7d_clean/run_20260410_171624/analysis/15min_windows/policy_provider_shares_15min_focus.png",
            caption=(
                "Alibaba is selected only in a subset of windows; it is not a "
                "uniform preference over the whole run."
            ),
        ),
        FigureSpec(
            title="Figure 5. Qwen3 Provider TTFT Percentiles (lp_mix only)",
            path=ROOT
            / "experiment/results/phase5_qwen3_7d_clean/run_20260410_171624/analysis/15min_windows/provider_ttft_percentiles_15min_lp_mix_only.png",
            caption=(
                "In Alibaba-heavy windows, WandB median remains reasonable, but "
                "WandB tail degrades noticeably."
            ),
        ),
        FigureSpec(
            title="Figure 6. LP Reconstruction: Alibaba Weight vs Realized Share",
            path=ROOT
            / "experiment/results/phase5_qwen3_7d_clean/run_20260410_171624/analysis/lp_reconstruction/alibaba_weight_vs_share.png",
            caption=(
                "The LP itself already assigns substantial mass to Alibaba in "
                "Alibaba-heavy windows; this is not purely a SWRR artifact."
            ),
        ),
        FigureSpec(
            title="Figure 7. LP Reconstruction: CDF and P99 Focus",
            path=ROOT
            / "experiment/results/phase5_qwen3_7d_clean/run_20260410_171624/analysis/lp_reconstruction/alibaba_cdf_focus.png",
            caption=(
                "When WandB F(SLO=2s) falls below the hard 0.99 target, Alibaba "
                "often stays closer to the target and receives large LP weight."
            ),
        ),
        FigureSpec(
            title="Figure 8. Counterfactual (Qwen3 no-WandB)",
            path=ROOT
            / "experiment/results/counterfactual_qwen3_no_wandb_v2/counterfactual_slo_p99.png",
            caption=(
                "Without WandB, smart_hedge shows meaningful tail improvement, "
                "but lp_mix still loses to the strongest fixed baseline."
            ),
        ),
        FigureSpec(
            title="Figure 9. Counterfactual (MiniMax no-Inceptron)",
            path=ROOT
            / "experiment/results/counterfactual_minimax_no_inceptron_v2/counterfactual_slo_p99.png",
            caption=(
                "Without Inceptron, strongest fixed baseline remains much better; "
                "current LP/CDF formulation is still not robust."
            ),
        ),
    ]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def _add_figure(doc: Document, spec: FigureSpec) -> None:
    _add_heading(doc, spec.title, level=2)
    doc.add_picture(str(spec.path), width=Inches(spec.width_in))
    caption = doc.add_paragraph(spec.caption)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True


def _build_doc() -> Document:
    online_summary = _load_online_summary(LATEST_V2_RUN)
    qwen_counterfactual = _load_counterfactual_rows(QWEN_COUNTERFACTUAL)
    minimax_counterfactual = _load_counterfactual_rows(MINIMAX_COUNTERFACTUAL)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RouteWise Meeting Update for Juncheng")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Generated on 2026-04-14 from the latest local analysis artifacts.")

    _add_heading(doc, "Executive Summary", level=1)
    _add_bullets(
        doc,
        [
            "Production ACF and live-run ACF both support the same direction: P50 is a cleaner online selection signal than P99.",
            "The Alibaba-heavy windows are now largely explained: the LP itself assigns Alibaba substantial mass when WandB tail degrades and F(SLO=2s) falls below target.",
            "No-dominant counterfactual experiments show that the problem is not only a single dominant-provider artifact; the current LP/CDF formulation is itself fragile.",
            "The first V2 online run shows that P50-based primary selection fixes LP over-diversification, but the tested window is a dominant-provider regime, so strongest fixed baseline still wins.",
        ],
    )

    _add_heading(doc, "1. ACF: P50 vs P99", level=1)
    doc.add_paragraph(
        "Goal: test whether recent probing is more informative for P50 or P99. "
        "We used both production e2e probes (main evidence) and a live Qwen3 run "
        "(supporting evidence)."
    )
    for spec in _get_figures()[:3]:
        _add_figure(doc, spec)
    _add_bullets(
        doc,
        [
            "Main conclusion: P50 has stronger short-term temporal correlation than P99 for most provider/model pairs.",
            "Interpretation: probing is better suited for primary-provider selection via P50 than for direct tail prediction via P99.",
            "System implication: probe for P50; protect P99 with runtime hedging.",
        ],
    )

    _add_heading(doc, "2. Why lp_mix Routes So Much Traffic to Alibaba", level=1)
    doc.add_paragraph(
        "Goal: explain both why Alibaba is selected and why the share can become so large. "
        "We combined 15-minute window analysis with an incremental replay of the actual "
        "OnlineLatencyRouter logic."
    )
    for spec in _get_figures()[3:7]:
        _add_figure(doc, spec)
    _add_bullets(
        doc,
        [
            "Alibaba is not selected uniformly; it is elevated only in a subset of windows.",
            "In those windows, WandB median remains competitive, but WandB tail degrades materially.",
            "The LP itself already assigns Alibaba large raw weight in those windows, so the effect is not purely a downstream SWRR artifact.",
            "Mechanism: when WandB F(SLO=2s) falls below the hard 0.99 target and Alibaba stays closer to target, the LP objective pushes mass toward Alibaba.",
        ],
    )

    _add_heading(doc, "3. No-Dominant Counterfactual Diagnosis", level=1)
    doc.add_paragraph(
        "Goal: test whether the failure mode is only caused by a dominant provider, or whether "
        "the LP/CDF formulation is itself not robust."
    )
    for spec in _get_figures()[7:]:
        _add_figure(doc, spec)

    _add_heading(doc, "Qwen3 no-WandB Summary", level=2)
    _add_table(
        doc,
        ["Policy", "SLO Viol", "95% CI", "P99 (ms)", "Mean Cost", "Hedge Rate"],
        [
            [
                row["policy"],
                row["slo"],
                row["slo_ci"],
                row["p99"],
                row["cost"],
                row["hedge_rate"],
            ]
            for row in qwen_counterfactual
        ],
    )
    _add_heading(doc, "MiniMax no-Inceptron Summary", level=2)
    _add_table(
        doc,
        ["Policy", "SLO Viol", "95% CI", "P99 (ms)", "Mean Cost", "Hedge Rate"],
        [
            [
                row["policy"],
                row["slo"],
                row["slo_ci"],
                row["p99"],
                row["cost"],
                row["hedge_rate"],
            ]
            for row in minimax_counterfactual
        ],
    )
    _add_bullets(
        doc,
        [
            "Qwen3 no-WandB: lp_mix still loses to the strongest fixed baseline; smart_hedge helps tail more than mixing.",
            "MiniMax no-Inceptron: strongest fixed baseline remains much better; this is not only a WandB artifact.",
            "Interpretation: the current LP/CDF formulation is not robust, even after removing the dominant provider.",
        ],
    )

    _add_heading(doc, "4. First V2 Online Run (Qwen3, 90 minutes)", level=1)
    doc.add_paragraph(
        "This run used trace replay with real prompts and the real default provider set. "
        "The purpose was to test whether replacing LP mixing with P50-based primary selection "
        "fixes the main pathology."
    )
    _add_table(
        doc,
        ["Policy", "Requests", "SLO Viol", "P50 (ms)", "P99 (ms)", "Total Cost"],
        [
            [
                row["policy"],
                row["n"],
                row["slo"],
                row["p50"],
                row["p99"],
                row["total_cost"],
            ]
            for row in online_summary
        ],
    )
    _add_bullets(
        doc,
        [
            "V2 substantially improves over lp_mix, confirming that P50-based primary selection fixes LP over-diversification.",
            "However, this window is a clear dominant-provider regime: cheapest_fixed and sort_latency both effectively collapse to WandB.",
            "Therefore this run validates collapse behavior, but does not yet show a decisive V2 advantage over the strongest fixed baseline.",
        ],
    )

    _add_heading(doc, "5. Current Interpretation", level=1)
    _add_bullets(
        doc,
        [
            "P50 is the right signal for online provider selection.",
            "Tail should be handled by hedging, not by hard LP/CDF mixing.",
            "The old LP/CDF router overreacts to short-lived tail fluctuations and pushes too much mass away from the dominant provider.",
            "V2 fixes the main LP pathology, but we still need a non-dominant or weak-dominant online run to test whether V2 can outperform strong baselines in the regime where a smart router should matter.",
        ],
    )

    _add_heading(doc, "6. Proposed Next Step", level=1)
    _add_bullets(
        doc,
        [
            "Run V2 in a non-dominant or weak-dominant regime, either by choosing a different model or by selecting a Qwen3 window where WandB is not overwhelmingly dominant.",
            "Keep the setup minimal: same probing cadence, same backup selection, same economic hedging, so the result remains attributable to the primary-selection rule.",
            "Use the next short online run to decide whether the paper story should shift from LP mixing to probe-for-P50 plus hedge-for-P99.",
        ],
    )

    return doc


def _build_markdown() -> str:
    online_summary = _load_online_summary(LATEST_V2_RUN)
    qwen_counterfactual = _load_counterfactual_rows(QWEN_COUNTERFACTUAL)
    minimax_counterfactual = _load_counterfactual_rows(MINIMAX_COUNTERFACTUAL)
    figures = _get_figures()

    lines = [
        "# RouteWise Meeting Update for Juncheng",
        "",
        "Generated on 2026-04-14 from the latest local analysis artifacts.",
        "",
        "## Executive Summary",
        "- Production ACF and live-run ACF both support the same direction: P50 is a cleaner online selection signal than P99.",
        "- The Alibaba-heavy windows are now largely explained: the LP itself assigns Alibaba substantial mass when WandB tail degrades and F(SLO=2s) falls below target.",
        "- No-dominant counterfactual experiments show that the problem is not only a single dominant-provider artifact; the current LP/CDF formulation is itself fragile.",
        "- The first V2 online run shows that P50-based primary selection fixes LP over-diversification, but the tested window is a dominant-provider regime, so strongest fixed baseline still wins.",
        "",
        "## 1. ACF: P50 vs P99",
        "Goal: test whether recent probing is more informative for P50 or P99.",
        "",
    ]
    for spec in figures[:3]:
        lines.extend([f"### {spec.title}", f"![{spec.title}]({spec.path})", spec.caption, ""])

    lines.extend(
        [
            "## 2. Why lp_mix Routes So Much Traffic to Alibaba",
            "Goal: explain both why Alibaba is selected and why the share can become so large.",
            "",
        ]
    )
    for spec in figures[3:7]:
        lines.extend([f"### {spec.title}", f"![{spec.title}]({spec.path})", spec.caption, ""])

    lines.extend(["## 3. No-Dominant Counterfactual Diagnosis", ""])
    for spec in figures[7:]:
        lines.extend([f"### {spec.title}", f"![{spec.title}]({spec.path})", spec.caption, ""])

    lines.extend(["### Qwen3 no-WandB Summary", ""])
    lines.append("| Policy | SLO Viol | 95% CI | P99 (ms) | Mean Cost | Hedge Rate |")
    lines.append("|---|---:|---|---:|---:|---:|")
    for row in qwen_counterfactual:
        lines.append(
            f"| {row['policy']} | {row['slo']} | {row['slo_ci']} | {row['p99']} | {row['cost']} | {row['hedge_rate']} |"
        )
    lines.extend(["", "### MiniMax no-Inceptron Summary", ""])
    lines.append("| Policy | SLO Viol | 95% CI | P99 (ms) | Mean Cost | Hedge Rate |")
    lines.append("|---|---:|---|---:|---:|---:|")
    for row in minimax_counterfactual:
        lines.append(
            f"| {row['policy']} | {row['slo']} | {row['slo_ci']} | {row['p99']} | {row['cost']} | {row['hedge_rate']} |"
        )

    lines.extend(["", "## 4. First V2 Online Run (Qwen3, 90 minutes)", ""])
    lines.append("| Policy | Requests | SLO Viol | P50 (ms) | P99 (ms) | Total Cost |")
    lines.append("|---|---:|---:|---:|---:|---:|")
    for row in online_summary:
        lines.append(
            f"| {row['policy']} | {row['n']} | {row['slo']} | {row['p50']} | {row['p99']} | {row['total_cost']} |"
        )

    lines.extend(
        [
            "",
            "## 5. Current Interpretation",
            "- P50 is the right signal for online provider selection.",
            "- Tail should be handled by hedging, not by hard LP/CDF mixing.",
            "- The old LP/CDF router overreacts to short-lived tail fluctuations and pushes too much mass away from the dominant provider.",
            "- V2 fixes the main LP pathology, but we still need a non-dominant or weak-dominant online run to test whether V2 can outperform strong baselines in the regime where a smart router should matter.",
            "",
            "## 6. Proposed Next Step",
            "- Run V2 in a non-dominant or weak-dominant regime.",
            "- Keep the setup minimal so the result stays attributable to the primary-selection rule.",
            "- Use the next short online run to decide whether the paper story should shift from LP mixing to probe-for-P50 plus hedge-for-P99.",
            "",
        ]
    )
    return "\n".join(lines)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = _build_doc()
    doc.save(DOCX_OUTPUT)
    MD_OUTPUT.write_text(_build_markdown())
    print(f"Wrote {DOCX_OUTPUT}")
    print(f"Wrote {MD_OUTPUT}")


if __name__ == "__main__":
    main()
